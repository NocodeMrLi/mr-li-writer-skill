import contextlib
import html
import importlib.util
import io
import json
import pathlib
import re
import subprocess
import sys
import tempfile
import unittest
from unittest import mock


ROOT = pathlib.Path(__file__).resolve().parents[1]


def load_module(name, path):
    spec = importlib.util.spec_from_file_location(name, ROOT / path)
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


build_gzh = load_module("build_gzh_html", "scripts/build_gzh_html.py")
build_html = load_module("build_html", "scripts/build_html.py")
extract_seed = load_module("extract_seed", "scripts/extract_seed.py")
lint_article = load_module("lint_article", "scripts/lint_article.py")


SAMPLE_MD = """# 测试文章

开头说明。

## 第一部分：先判断问题

正文一。

## 第二部分：再给方法

正文二。

## 第三部分：完成交付

正文三。
"""

TABLE_MD = """# 报名状态

| 省份 | 报名开始 | 报名截止 | 备注 |
| --- | --- | --- | --- |
| 广东 | 8月17日 9:00 | 8月25日 17:00 | 今天需要完成缴费并确认报名状态 |
| 辽宁 | 8月19日 8:30 | 8月25日 16:00 | 不含大连，大连单独安排 |
"""


class DeliveryProtocolTests(unittest.TestCase):
    def write_task_state(self, directory, **overrides):
        base = {
            "original_prompt": "写一篇公众号文章，主题是测试文章",
            "platform": {"value": "公众号", "confirmed": True, "source": "user", "user_quote": "公众号"},
            "content_goal": {"value": "普通传播", "confirmed": True, "source": "user", "user_quote": "普通传播"},
            "writing_direction": {"value": "实用指南", "confirmed": True, "source": "user", "user_quote": "写成实用指南"},
            "delivery_style": {"value": "moyu-green", "confirmed": True, "source": "user", "user_quote": "摸鱼绿"},
        }
        base.update(overrides)
        path = pathlib.Path(directory) / "task-state.json"
        path.write_text(json.dumps(base, ensure_ascii=False), encoding="utf-8")
        return path

    def test_task_intake_blocks_resume_when_required_confirmations_are_missing(self):
        validator = ROOT / "scripts/validate_task_intake.py"
        with tempfile.TemporaryDirectory() as tmp:
            state = pathlib.Path(tmp) / "task-state.json"
            state.write_text(
                json.dumps(
                    {
                        "platform": {
                            "value": "公众号",
                            "confirmed": True,
                            "source": "user",
                            "user_quote": "公众号",
                        }
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            result = subprocess.run(
                [
                    sys.executable,
                    str(validator),
                    str(state),
                    "--phase",
                    "resume",
                    "--last-user",
                    "继续完成任务",
                ],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("不能把“继续完成任务”视为确认", result.stdout)
            self.assertIn("内容目标", result.stdout)
            self.assertIn("创作方向", result.stdout)
            self.assertIn("平台交付样式", result.stdout)

    def test_task_intake_rejects_model_inferred_confirmations(self):
        validator = ROOT / "scripts/validate_task_intake.py"
        with tempfile.TemporaryDirectory() as tmp:
            state = self.write_task_state(
                tmp,
                delivery_style={
                    "value": "graphite-minimal",
                    "confirmed": True,
                    "source": "model",
                    "user_quote": "模型推荐石墨极简",
                },
            )
            result = subprocess.run(
                [sys.executable, str(validator), str(state), "--phase", "layout"],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("平台交付样式", result.stdout)
            self.assertIn("不得用模型推断", result.stdout)

    def test_task_intake_rejects_cross_platform_delivery_style(self):
        validator = ROOT / "scripts/validate_task_intake.py"
        with tempfile.TemporaryDirectory() as tmp:
            state = self.write_task_state(
                tmp,
                platform={"value": "知乎", "confirmed": True, "source": "user", "user_quote": "知乎"},
                delivery_style={"value": "graphite-minimal", "confirmed": True, "source": "user", "user_quote": "石墨极简风"},
            )
            result = subprocess.run(
                [sys.executable, str(validator), str(state), "--phase", "layout", "--platform", "知乎"],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("不得把公众号主题或其他平台样式跨平台套用", result.stdout)
            self.assertIn("回答、专栏", result.stdout)

    def test_task_intake_accepts_user_confirmed_or_authorized_auto_match(self):
        validator = ROOT / "scripts/validate_task_intake.py"
        with tempfile.TemporaryDirectory() as tmp:
            state = self.write_task_state(
                tmp,
                delivery_style={
                    "value": "auto",
                    "confirmed": True,
                    "source": "auto_authorized",
                    "user_quote": "不用问，自动匹配",
                },
            )
            result = subprocess.run(
                [sys.executable, str(validator), str(state), "--phase", "layout"],
                capture_output=True,
                text=True,
            )
            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)

    def test_task_intake_rejects_memory_or_standing_instruction_as_authorization(self):
        validator = ROOT / "scripts/validate_task_intake.py"
        with tempfile.TemporaryDirectory() as tmp:
            state = self.write_task_state(
                tmp,
                delivery_style={
                    "value": "auto",
                    "confirmed": True,
                    "source": "auto_authorized",
                    "user_quote": "根据长期偏好：意图明确时直接执行、不询问、排版自行决定",
                },
            )
            result = subprocess.run(
                [sys.executable, str(validator), str(state), "--phase", "layout"],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("长期记忆", result.stdout)

    def test_task_intake_from_prompt_blocks_new_task_without_confirmations(self):
        validator = ROOT / "scripts/validate_task_intake.py"
        result = subprocess.run(
            [
                sys.executable,
                str(validator),
                "--from-prompt",
                "Meta AI 原生组织转型受挫，原文：https://example.com/article",
                "--phase",
                "task-list",
            ],
            capture_output=True,
            text=True,
        )
        self.assertNotEqual(result.returncode, 0)
        self.assertIn("新任务尚未形成任务状态", result.stdout)
        self.assertIn("发布平台", result.stdout)
        self.assertIn("内容目标", result.stdout)
        self.assertIn("创作方向", result.stdout)
        self.assertIn("平台交付样式", result.stdout)

    def test_task_intake_from_prompt_does_not_auto_authorize_from_memory_words(self):
        validator = ROOT / "scripts/validate_task_intake.py"
        result = subprocess.run(
            [
                sys.executable,
                str(validator),
                "--from-prompt",
                "根据长期偏好直接处理，写成公众号并自动匹配",
                "--phase",
                "task-list",
            ],
            capture_output=True,
            text=True,
        )
        self.assertNotEqual(result.returncode, 0)
        self.assertIn("发布平台", result.stdout)
        self.assertIn("当前平台可选交付样式", result.stdout)

    def test_task_intake_from_prompt_rejects_craft_mode_memory_authorization(self):
        validator = ROOT / "scripts/validate_task_intake.py"
        result = subprocess.run(
            [
                sys.executable,
                str(validator),
                "--from-prompt",
                "按 Craft mode 和用户偏好直接处理，这次排版由 AI 自行决定并自动匹配。",
                "--phase",
                "task-list",
            ],
            capture_output=True,
            text=True,
        )
        self.assertNotEqual(result.returncode, 0)
        self.assertIn("发布平台", result.stdout)
        self.assertIn("平台交付样式", result.stdout)
        self.assertIn("当前平台可选交付样式", result.stdout)

    def test_task_intake_rejects_platform_value_mixed_with_layout_theme(self):
        validator = ROOT / "scripts/validate_task_intake.py"
        with tempfile.TemporaryDirectory() as tmp:
            state = self.write_task_state(
                tmp,
                platform={
                    "value": "公众号（橄榄手记）",
                    "confirmed": True,
                    "source": "user",
                    "user_quote": "公众号",
                },
                delivery_style={
                    "value": "橄榄手记",
                    "confirmed": True,
                    "source": "user",
                    "user_quote": "橄榄手记",
                },
            )
            result = subprocess.run(
                [sys.executable, str(validator), str(state), "--phase", "layout", "--platform", "公众号"],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("发布平台不能混入排版主题", result.stdout)

    def test_task_intake_emit_question_card_separates_platform_and_layout_theme(self):
        validator = ROOT / "scripts/validate_task_intake.py"
        result = subprocess.run(
            [
                sys.executable,
                str(validator),
                "--from-prompt",
                "这篇 Gemini 3.5 Transcribe 的文章发到哪个平台？",
                "--phase",
                "task-list",
                "--emit-question-card",
            ],
            capture_output=True,
            text=True,
        )
        self.assertNotEqual(result.returncode, 0)
        self.assertIn("请先确认以下信息", result.stdout)
        self.assertIn("1. 发布平台", result.stdout)
        self.assertIn("公众号 / 小红书 / 知乎 / 官网/网页 / 个人博客", result.stdout)
        self.assertIn("2. 内容目标", result.stdout)
        self.assertIn("3. 创作方向", result.stdout)
        self.assertIn("4. 平台交付样式", result.stdout)
        self.assertIn("最推荐", result.stdout)
        self.assertIn("次推荐", result.stdout)
        self.assertIn("补充说明/自行输入", result.stdout)
        self.assertIn("公众号排版主题不是发布平台", result.stdout)
        self.assertNotIn("公众号（橄榄手记）", result.stdout)

    def test_task_intake_from_prompt_always_outputs_standard_question_card(self):
        validator = ROOT / "scripts/validate_task_intake.py"
        result = subprocess.run(
            [
                sys.executable,
                str(validator),
                "--from-prompt",
                "原文：https://example.com/article，帮我写一篇文章",
                "--phase",
                "task-list",
            ],
            capture_output=True,
            text=True,
        )
        self.assertNotEqual(result.returncode, 0)
        self.assertIn("请先确认以下信息", result.stdout)
        self.assertIn("1. 发布平台", result.stdout)
        self.assertIn("2. 内容目标", result.stdout)
        self.assertIn("3. 创作方向", result.stdout)
        self.assertIn("4. 平台交付样式", result.stdout)
        self.assertIn("补充说明/自行输入", result.stdout)

    def test_task_intake_from_prompt_does_not_treat_title_auto_match_as_global_authorization(self):
        validator = ROOT / "scripts/validate_task_intake.py"
        result = subprocess.run(
            [
                sys.executable,
                str(validator),
                "--from-prompt",
                "帮我写公众号文章，标题可以自动匹配，其他先问我。",
                "--phase",
                "task-list",
            ],
            capture_output=True,
            text=True,
        )
        self.assertNotEqual(result.returncode, 0)
        self.assertIn("内容目标", result.stdout)
        self.assertIn("创作方向", result.stdout)
        self.assertIn("平台交付样式", result.stdout)

    def test_task_intake_from_prompt_does_not_treat_theme_auto_match_as_global_authorization(self):
        validator = ROOT / "scripts/validate_task_intake.py"
        result = subprocess.run(
            [
                sys.executable,
                str(validator),
                "--from-prompt",
                "帮我写公众号文章，公众号主题自动匹配，其他先问我。",
                "--phase",
                "task-list",
            ],
            capture_output=True,
            text=True,
        )
        self.assertNotEqual(result.returncode, 0)
        self.assertIn("内容目标", result.stdout)
        self.assertIn("创作方向", result.stdout)
        self.assertIn("平台交付样式", result.stdout)

    def test_task_intake_from_prompt_does_not_treat_local_layout_instruction_as_global_authorization(self):
        validator = ROOT / "scripts/validate_task_intake.py"
        result = subprocess.run(
            [
                sys.executable,
                str(validator),
                "--from-prompt",
                "帮我写公众号文章，这篇文章直接排不用管目录，其他先问我。",
                "--phase",
                "task-list",
            ],
            capture_output=True,
            text=True,
        )
        self.assertNotEqual(result.returncode, 0)
        self.assertIn("内容目标", result.stdout)
        self.assertIn("创作方向", result.stdout)
        self.assertIn("平台交付样式", result.stdout)

    def test_task_intake_from_prompt_does_not_treat_title_you_decide_as_global_authorization(self):
        validator = ROOT / "scripts/validate_task_intake.py"
        result = subprocess.run(
            [
                sys.executable,
                str(validator),
                "--from-prompt",
                "帮我写公众号文章，标题你看着办。",
                "--phase",
                "task-list",
            ],
            capture_output=True,
            text=True,
        )
        self.assertNotEqual(result.returncode, 0)
        self.assertIn("内容目标", result.stdout)
        self.assertIn("创作方向", result.stdout)
        self.assertIn("平台交付样式", result.stdout)

    def test_task_intake_from_prompt_does_not_treat_theme_you_decide_as_global_authorization(self):
        validator = ROOT / "scripts/validate_task_intake.py"
        result = subprocess.run(
            [
                sys.executable,
                str(validator),
                "--from-prompt",
                "帮我写公众号文章，主题你看着办。",
                "--phase",
                "task-list",
            ],
            capture_output=True,
            text=True,
        )
        self.assertNotEqual(result.returncode, 0)
        self.assertIn("内容目标", result.stdout)
        self.assertIn("创作方向", result.stdout)
        self.assertIn("平台交付样式", result.stdout)

    def test_task_intake_from_prompt_does_not_treat_title_no_ask_as_global_authorization(self):
        validator = ROOT / "scripts/validate_task_intake.py"
        result = subprocess.run(
            [
                sys.executable,
                str(validator),
                "--from-prompt",
                "帮我写公众号文章，标题不用问。",
                "--phase",
                "task-list",
            ],
            capture_output=True,
            text=True,
        )
        self.assertNotEqual(result.returncode, 0)
        self.assertIn("内容目标", result.stdout)
        self.assertIn("创作方向", result.stdout)
        self.assertIn("平台交付样式", result.stdout)

    def test_task_intake_from_prompt_does_not_treat_theme_direct_process_as_global_authorization(self):
        validator = ROOT / "scripts/validate_task_intake.py"
        result = subprocess.run(
            [
                sys.executable,
                str(validator),
                "--from-prompt",
                "帮我写公众号文章，主题直接处理。",
                "--phase",
                "task-list",
            ],
            capture_output=True,
            text=True,
        )
        self.assertNotEqual(result.returncode, 0)
        self.assertIn("内容目标", result.stdout)
        self.assertIn("创作方向", result.stdout)
        self.assertIn("平台交付样式", result.stdout)

    def test_task_intake_from_prompt_does_not_treat_title_you_decide_ok_as_global_authorization(self):
        validator = ROOT / "scripts/validate_task_intake.py"
        result = subprocess.run(
            [
                sys.executable,
                str(validator),
                "--from-prompt",
                "帮我写公众号文章，标题你看着办就行。",
                "--phase",
                "task-list",
            ],
            capture_output=True,
            text=True,
        )
        self.assertNotEqual(result.returncode, 0)
        self.assertIn("内容目标", result.stdout)
        self.assertIn("创作方向", result.stdout)
        self.assertIn("平台交付样式", result.stdout)

    def test_task_intake_from_prompt_all_you_decide_is_global_authorization(self):
        validator = ROOT / "scripts/validate_task_intake.py"
        result = subprocess.run(
            [
                sys.executable,
                str(validator),
                "--from-prompt",
                "帮我写公众号文章，本次全部你看着办。",
                "--phase",
                "task-list",
            ],
            capture_output=True,
            text=True,
        )
        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)

    def test_task_intake_resume_detection_does_not_trigger_on_incidental_english_continue(self):
        validator = ROOT / "scripts/validate_task_intake.py"
        with tempfile.TemporaryDirectory() as tmp:
            state = pathlib.Path(tmp) / "task-state.json"
            state.write_text(
                json.dumps(
                    {
                        "original_prompt": "写一篇文章",
                        "platform": {
                            "value": "公众号",
                            "confirmed": True,
                            "source": "user",
                            "user_quote": "公众号",
                        },
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            result = subprocess.run(
                [
                    sys.executable,
                    str(validator),
                    str(state),
                    "--phase",
                    "draft",
                    "--last-user",
                    "please continue reading the source before writing",
                ],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertNotIn("不能把“继续完成任务”视为确认", result.stdout)

    def test_task_intake_resume_detection_covers_natural_chinese_variants(self):
        validator = load_module("validate_task_intake_resume_variants", "scripts/validate_task_intake.py")
        for text in (
            "继续吧",
            "继续把排版做完",
            "继续把正文写完",
            "继续完成剩下的排版",
            "继续完成剩余交付",
            "接着把交付做完",
        ):
            with self.subTest(text=text):
                self.assertTrue(validator.RESUME_RE.search(text))

    def test_task_intake_rejects_from_prompt_and_state_together(self):
        validator = ROOT / "scripts/validate_task_intake.py"
        with tempfile.TemporaryDirectory() as tmp:
            state = self.write_task_state(tmp)
            result = subprocess.run(
                [
                    sys.executable,
                    str(validator),
                    str(state),
                    "--from-prompt",
                    "不用问，自动匹配",
                    "--phase",
                    "task-list",
                ],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("不能同时使用", result.stdout + result.stderr)

    def test_task_intake_does_not_silently_disable_research_scope_gate(self):
        source = (ROOT / "scripts/validate_task_intake.py").read_text(encoding="utf-8")
        self.assertNotIn("validate_research_scope = None", source)

    def test_task_intake_reports_missing_research_scope_validator_without_traceback(self):
        source = (ROOT / "scripts/validate_task_intake.py").read_text(encoding="utf-8")
        with tempfile.TemporaryDirectory() as tmp:
            directory = pathlib.Path(tmp)
            script = directory / "validate_task_intake.py"
            script.write_text(source, encoding="utf-8")
            state = self.write_task_state(directory)
            result = subprocess.run(
                [sys.executable, str(script), str(state), "--phase", "draft"],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("资料搜集范围校验脚本", result.stdout)
            self.assertNotIn("Traceback", result.stdout + result.stderr)

    def test_research_scope_blocks_user_link_only_for_time_sensitive_hard_info(self):
        validator = ROOT / "scripts/validate_research_scope.py"
        with tempfile.TemporaryDirectory() as tmp:
            state = self.write_task_state(
                tmp,
                original_prompt="https://mp.weixin.qq.com/s/example 2026 下半年软考报名要求",
                research_scope={
                    "user_seed_sources": ["https://mp.weixin.qq.com/s/example"],
                    "user_sources_checked": True,
                    "requires_external_research": True,
                    "risk": "time_sensitive_hard_info",
                },
            )
            result = subprocess.run(
                [sys.executable, str(validator), str(state), "--phase", "draft"],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("种子资料不是完整资料搜集", result.stdout)
            self.assertIn("官方/权威来源", result.stdout)
            self.assertIn("最新核验", result.stdout)

    def test_research_scope_accepts_seed_plus_official_fresh_crosscheck(self):
        validator = ROOT / "scripts/validate_research_scope.py"
        with tempfile.TemporaryDirectory() as tmp:
            state = self.write_task_state(
                tmp,
                original_prompt="https://mp.weixin.qq.com/s/example 2026 下半年软考报名要求",
                research_scope={
                    "user_seed_sources": ["https://mp.weixin.qq.com/s/example"],
                    "user_sources_checked": True,
                    "external_search_done": True,
                    "official_sources_checked": True,
                    "freshness_checked": True,
                    "independent_crosscheck_checked": True,
                    "requires_external_research": True,
                    "risk": "time_sensitive_hard_info",
                    "source_mix": "用户种子资料 + 全国软考报名平台 + 各省软考办通知",
                    "source_entities": [
                        {
                            "name": "全国软考报名平台",
                            "role": "official",
                            "claim_scope": "2026 年软考报名安排",
                            "reader_visibility": "named",
                        }
                    ],
                },
            )
            result = subprocess.run(
                [sys.executable, str(validator), str(state), "--phase", "draft"],
                capture_output=True,
                text=True,
            )
            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)

    def test_research_scope_allows_explicit_seed_only_boundary(self):
        validator = ROOT / "scripts/validate_research_scope.py"
        with tempfile.TemporaryDirectory() as tmp:
            state = self.write_task_state(
                tmp,
                original_prompt="只基于我给的资料写，不再外查：https://mp.weixin.qq.com/s/example",
                research_scope={
                    "user_seed_sources": ["https://mp.weixin.qq.com/s/example"],
                    "user_sources_checked": True,
                    "requires_external_research": True,
                    "risk": "time_sensitive_hard_info",
                    "source_entities": [
                        {
                            "name": "用户提供资料",
                            "role": "unknown",
                            "claim_scope": "有限资料整理",
                            "reader_visibility": "anonymous",
                        }
                    ],
                },
                source_boundary={
                    "value": "只基于用户给的资料，不再外查",
                    "confirmed": True,
                    "source": "user",
                    "user_quote": "只基于我给的资料写，不再外查",
                },
            )
            result = subprocess.run(
                [sys.executable, str(validator), str(state), "--phase", "draft"],
                capture_output=True,
                text=True,
            )
            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)

    def test_research_scope_blocks_commercial_entity_named_as_authority(self):
        validator = ROOT / "scripts/validate_research_scope.py"
        with tempfile.TemporaryDirectory() as tmp:
            state = self.write_task_state(
                tmp,
                original_prompt="2026 年证书考试报名信息",
                research_scope={
                    "external_search_done": True,
                    "official_sources_checked": True,
                    "freshness_checked": True,
                    "independent_crosscheck_checked": True,
                    "requires_external_research": True,
                    "source_mix": "官方信息与第三方交叉核对",
                    "source_entities": [
                        {
                            "name": "星河学习",
                            "role": "commercial_interested",
                            "claim_scope": "考试城市汇总",
                            "reader_visibility": "named",
                        }
                    ],
                },
            )
            result = subprocess.run(
                [sys.executable, str(validator), str(state), "--phase", "draft"],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("不能在正文显名", result.stdout)

    def test_task_intake_requires_task_context_before_draft(self):
        validator = ROOT / "scripts/validate_task_intake.py"
        with tempfile.TemporaryDirectory() as tmp:
            state = self.write_task_state(tmp, original_prompt="")
            result = subprocess.run(
                [sys.executable, str(validator), str(state), "--phase", "draft"],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("original_prompt/topic/brief", result.stdout)
            self.assertIn("不能绕过资料搜集范围校验", result.stdout)

    def test_task_intake_invokes_research_scope_before_draft(self):
        validator = ROOT / "scripts/validate_task_intake.py"
        with tempfile.TemporaryDirectory() as tmp:
            state = self.write_task_state(
                tmp,
                original_prompt="https://mp.weixin.qq.com/s/example 2026 下半年软考报名要求",
                research_scope={
                    "user_seed_sources": ["https://mp.weixin.qq.com/s/example"],
                    "requires_external_research": True,
                    "risk": "time_sensitive_hard_info",
                },
            )
            result = subprocess.run(
                [sys.executable, str(validator), str(state), "--phase", "draft"],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("资料搜集", result.stdout)
            self.assertIn("不能只解析用户提供的链接", result.stdout)

    def test_documents_require_intake_check_before_fetching_or_task_list(self):
        skill = (ROOT / "SKILL.md").read_text(encoding="utf-8")
        agent = (ROOT / "agents/openai.yaml").read_text(encoding="utf-8")
        for text in (skill, agent):
            self.assertIn("validate_task_intake.py", text)
            self.assertIn("--from-prompt", text)
            self.assertIn("--emit-question-card", text)
        self.assertIn("抓取、检索、读取链接、生成任务列表", skill)
        self.assertIn("before fetching source links", agent)

    def test_documents_require_seed_source_expansion_protocol(self):
        skill = (ROOT / "SKILL.md").read_text(encoding="utf-8")
        research = (ROOT / "references/research-protocol.md").read_text(encoding="utf-8")
        readme = (ROOT / "README.md").read_text(encoding="utf-8")
        agent = (ROOT / "agents/openai.yaml").read_text(encoding="utf-8")
        for text in (skill, research, readme, agent):
            self.assertIn("种子资料", text)
            self.assertIn("validate_research_scope.py", text)
        for token in (
            "user_seed_sources",
            "external_search_done",
            "official_sources_checked",
            "freshness_checked",
            "source_mix",
            "source_entities",
        ):
            self.assertIn(token, skill + research + agent)
        self.assertIn("只基于我给的资料", skill + research + readme)
        self.assertIn("二创整理", skill + research + readme)

    def test_build_html_requires_task_state_before_formal_generation(self):
        builder = ROOT / "scripts/build_html.py"
        with tempfile.TemporaryDirectory() as tmp:
            directory = pathlib.Path(tmp)
            source = directory / "article-source.md"
            output = directory / "article.html"
            source.write_text(SAMPLE_MD, encoding="utf-8")
            missing_state = subprocess.run(
                [
                    sys.executable,
                    str(builder),
                    str(source),
                    "-o",
                    str(output),
                    "--platform",
                    "公众号",
                    "-t",
                    "moyu-green",
                    "--theme-confirmed",
                ],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(missing_state.returncode, 0)
            self.assertIn("--task-state", missing_state.stderr + missing_state.stdout)

            state = self.write_task_state(directory)
            complete = subprocess.run(
                [
                    sys.executable,
                    str(builder),
                    str(source),
                    "-o",
                    str(output),
                    "--platform",
                    "公众号",
                    "-t",
                    "moyu-green",
                    "--theme-confirmed",
                    "--task-state",
                    str(state),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(complete.returncode, 0, complete.stdout + complete.stderr)

    def test_build_html_requires_task_state_for_every_platform(self):
        builder = ROOT / "scripts/build_html.py"
        platforms = ("知乎", "小红书", "官网/网页", "个人博客")
        with tempfile.TemporaryDirectory() as tmp:
            directory = pathlib.Path(tmp)
            source = directory / "article-source.md"
            source.write_text(SAMPLE_MD, encoding="utf-8")
            for platform in platforms:
                output = directory / ("%s.html" % platform.replace("/", "-"))
                result = subprocess.run(
                    [
                        sys.executable,
                        str(builder),
                        str(source),
                        "-o",
                        str(output),
                        "--platform",
                        platform,
                    ],
                    capture_output=True,
                    text=True,
            )
            self.assertNotEqual(result.returncode, 0, platform)
            self.assertIn("--task-state", result.stderr + result.stdout)

    def test_non_wechat_html_deduplicates_page_title_and_markdown_h1(self):
        builder = ROOT / "scripts/build_html.py"
        title = "PMP 新考纲 12 月 5 日首考：第七版和第八版一张表对照，AI 题会怎么出"
        with tempfile.TemporaryDirectory() as tmp:
            directory = pathlib.Path(tmp)
            source = directory / "zhihu-source.md"
            output = directory / "zhihu.html"
            source.write_text("# %s\n\n正文第一段。\n" % title, encoding="utf-8")
            state = self.write_task_state(
                directory,
                platform={"value": "知乎", "confirmed": True, "source": "user", "user_quote": "知乎"},
                delivery_style={"value": "回答", "confirmed": True, "source": "user", "user_quote": "回答"},
            )
            result = subprocess.run(
                [
                    sys.executable,
                    str(builder),
                    str(source),
                    "-o",
                    str(output),
                    "--platform",
                    "知乎",
                    "--delivery-style",
                    "zhihu-answer",
                    "--title",
                    title,
                    "--task-state",
                    str(state),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
            html_text = output.read_text(encoding="utf-8")
            main = re.search(r'<main class="article" id="article">(.*?)</main>', html_text, re.S).group(1)
            body_text = html.unescape(re.sub(r"<[^>]+>", "", main))
            self.assertEqual(body_text.count(title), 1)
            self.assertEqual(len(re.findall(r"<h1", html_text)), 1)

    def test_build_html_rejects_theme_that_differs_from_task_state(self):
        builder = ROOT / "scripts/build_html.py"
        with tempfile.TemporaryDirectory() as tmp:
            directory = pathlib.Path(tmp)
            source = directory / "article-source.md"
            output = directory / "article.html"
            source.write_text(SAMPLE_MD, encoding="utf-8")
            state = self.write_task_state(
                directory,
                delivery_style={
                    "value": "moyu-green",
                    "confirmed": True,
                    "source": "user",
                    "user_quote": "摸鱼绿",
                },
            )
            result = subprocess.run(
                [
                    sys.executable,
                    str(builder),
                    str(source),
                    "-o",
                    str(output),
                    "--platform",
                    "公众号",
                    "-t",
                    "graphite-minimal",
                    "--theme-confirmed",
                    "--task-state",
                    str(state),
                ],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("命令交付样式与任务状态不一致", result.stdout)

    def test_delivery_bundle_requires_task_state(self):
        validator = ROOT / "scripts/validate_delivery_bundle.py"
        with tempfile.TemporaryDirectory() as tmp:
            delivery = pathlib.Path(tmp)
            (delivery / "title-strategy.md").write_text(
                "# 标题策略\n\n## 主标题\n测试标题\n\n## 备选标题\n备选一、备选二。",
                encoding="utf-8",
            )
            (delivery / "article-source.md").write_text("# 正文\n\n内容。", encoding="utf-8")
            (delivery / "article.html").write_text("<main>正文</main>", encoding="utf-8")
            (delivery / "article-preview.html").write_text(
                "<button>复制正文</button><script>function copyArticle(){}</script>",
                encoding="utf-8",
            )
            missing_state = subprocess.run(
                [sys.executable, str(validator), str(delivery), "--platform", "公众号"],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(missing_state.returncode, 0)
            self.assertIn("--task-state", missing_state.stdout + missing_state.stderr)

            state = self.write_task_state(delivery)
            complete = subprocess.run(
                [
                    sys.executable,
                    str(validator),
                    str(delivery),
                    "--platform",
                    "公众号",
                    "--task-state",
                    str(state),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(complete.returncode, 0, complete.stdout + complete.stderr)

    def test_wrap_gzh_preview_requires_task_state(self):
        wrapper = ROOT / "scripts/wrap_gzh_preview.py"
        with tempfile.TemporaryDirectory() as tmp:
            directory = pathlib.Path(tmp)
            source = directory / "article.html"
            output = directory / "article-preview.html"
            source.write_text("<section><p><span leaf=\"\">正文</span></p></section>", encoding="utf-8")

            missing_state = subprocess.run(
                [sys.executable, str(wrapper), str(source), str(output)],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(missing_state.returncode, 0)
            self.assertIn("--task-state", missing_state.stderr + missing_state.stdout)

            state = self.write_task_state(directory)
            complete = subprocess.run(
                [sys.executable, str(wrapper), str(source), str(output), "--task-state", str(state)],
                capture_output=True,
                text=True,
            )
            self.assertEqual(complete.returncode, 0, complete.stdout + complete.stderr)
            self.assertTrue(output.is_file())

    def test_wrap_gzh_preview_help_uses_current_script_name(self):
        wrapper = ROOT / "scripts/wrap_gzh_preview.py"
        result = subprocess.run(
            [sys.executable, str(wrapper), "--help"],
            capture_output=True,
            text=True,
        )
        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        self.assertIn("wrap_gzh_preview.py", result.stdout)

    def test_docx_embedded_image_without_ocr_keeps_extracted_text(self):
        try:
            import docx
            from PIL import Image
        except ImportError as exc:
            self.skipTest("DOCX image fixture dependencies unavailable: %s" % exc)
        with tempfile.TemporaryDirectory() as tmp:
            directory = pathlib.Path(tmp)
            image_path = directory / "tiny.png"
            Image.new("RGB", (8, 8), "white").save(image_path)
            doc_path = directory / "seed.docx"
            document = docx.Document()
            document.add_paragraph("这段 DOCX 正文应该被保留。")
            document.add_picture(str(image_path))
            document.save(doc_path)
            with mock.patch.object(extract_seed, "ocr_image", side_effect=SystemExit(2)):
                text = extract_seed.extract_docx(str(doc_path))
        self.assertIn("这段 DOCX 正文应该被保留", text)
        self.assertIn("OCR", text)
        self.assertIn("跳过", text)

    def test_skill_requires_native_file_attachments(self):
        text = (ROOT / "SKILL.md").read_text(encoding="utf-8")
        self.assertIn("原生附件", text)
        self.assertNotIn("`[用途名](绝对路径) — 一句话用途`", text)

    def test_research_protocol_covers_commercial_source_conflicts(self):
        text = (ROOT / "references/research-protocol.md").read_text(encoding="utf-8")
        self.assertIn("直接商业利益", text)
        self.assertIn("相关机构公开汇总", text)

    def test_wechat_delivery_requires_and_validates_four_artifacts(self):
        validator = ROOT / "scripts/validate_delivery_bundle.py"
        protocol = (ROOT / "references/delivery-protocol.md").read_text(encoding="utf-8")
        self.assertTrue(validator.exists(), "缺少公众号四件套交付校验器")
        for role in ("标题策略", "正文 Markdown", "公众号正文 HTML", "复制预览 HTML"):
            self.assertIn(role, protocol)

        with tempfile.TemporaryDirectory() as tmp:
            delivery = pathlib.Path(tmp)
            state = self.write_task_state(delivery)
            (delivery / "article-source.md").write_text("# 正文\n\n内容。", encoding="utf-8")
            (delivery / "article.html").write_text("<section>正文</section>", encoding="utf-8")
            (delivery / "article-preview.html").write_text(
                "<button>复制到公众号</button><script>function gzhCopy(){}</script>",
                encoding="utf-8",
            )
            missing = subprocess.run(
                [sys.executable, str(validator), str(delivery), "--platform", "公众号", "--task-state", str(state)],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(missing.returncode, 0)

            (delivery / "title-strategy.md").write_text(
                "# 标题策略\n\n## 主标题\n测试标题\n\n## 备选标题\n备选一、备选二。",
                encoding="utf-8",
            )
            complete = subprocess.run(
                [sys.executable, str(validator), str(delivery), "--platform", "公众号", "--task-state", str(state)],
                capture_output=True,
                text=True,
            )
            self.assertEqual(complete.returncode, 0, complete.stdout + complete.stderr)

    def test_agent_entrypoint_repeats_wechat_delivery_guards(self):
        text = (ROOT / "agents/openai.yaml").read_text(encoding="utf-8")
        self.assertIn("four real", text)
        self.assertIn("semantic", text)
        self.assertIn("validate_delivery_bundle.py", text)
        self.assertIn("--auto-theme-ok", text)
        self.assertIn("--theme-confirmed", text)
        self.assertIn("If the publishing platform is missing", text)
        self.assertIn("3-5 concrete options", text)

    def test_resume_requests_must_recheck_required_confirmations(self):
        skill = (ROOT / "SKILL.md").read_text(encoding="utf-8")
        platform_protocol = (ROOT / "references/platform-native-protocol.md").read_text(encoding="utf-8")
        for text in (skill, platform_protocol):
            self.assertIn("恢复任务防误判", text)
            self.assertIn("继续完成任务", text)
            self.assertIn("不得把“继续”理解为用户已确认缺失项", text)
            self.assertIn("发布平台、内容目标、创作方向、平台交付样式", text)

    def test_platform_delivery_style_confirmation_has_no_interpretation_gap(self):
        skill = (ROOT / "SKILL.md").read_text(encoding="utf-8")
        platform_protocol = (ROOT / "references/platform-native-protocol.md").read_text(encoding="utf-8")
        agent = (ROOT / "agents/openai.yaml").read_text(encoding="utf-8")
        for text in (skill, platform_protocol):
            self.assertIn("平台交付样式硬确认", text)
            self.assertIn("写正文、生成任务列表、创建正文文件、生成 HTML 或交付文件之前", text)
            self.assertIn("必须完整展示该平台的全部可选交付样式", text)
            self.assertIn("不得只展示推荐项", text)
            self.assertIn("公众号排版主题", text)
            for option in ("摸鱼绿", "红白色系", "石墨极简风", "留白禅意风", "摸鱼票据风", "橄榄手记", "自动匹配"):
                self.assertIn(option, text)
        for phrase in (
            "confirm the platform delivery style before drafting",
            "show the complete option set",
            "not only recommended options",
            "all six WeChat layout themes",
        ):
            self.assertIn(phrase, agent)

    def test_required_and_conditional_questions_are_hard_gates(self):
        skill = (ROOT / "SKILL.md").read_text(encoding="utf-8")
        platform_protocol = (ROOT / "references/platform-native-protocol.md").read_text(encoding="utf-8")
        agent = (ROOT / "agents/openai.yaml").read_text(encoding="utf-8")
        readme = (ROOT / "README.md").read_text(encoding="utf-8")
        for text in (skill, platform_protocol):
            self.assertIn("必问项与条件触发项硬门禁", text)
            self.assertIn("必问项缺失时不得写正文、生成任务列表、创建文件、排版或交付", text)
            self.assertIn("条件触发项一旦被触发，就临时升级为本任务的必确字段", text)
            self.assertIn("不得用模型推断、历史默认、平台默认或推荐项替代用户确认", text)
            for item in ("发布平台", "内容目标", "创作方向", "平台交付样式"):
                self.assertIn(item, text)
            for item in ("目标读者/阅读场景", "立场边界", "时效口径", "来源边界", "交付格式", "篇幅深度"):
                self.assertIn(item, text)
        for phrase in (
            "Required fields and triggered conditional fields are hard gates",
            "Do not draft, create tasks, create files, layout, or deliver",
            "Triggered conditional fields become required for that task",
        ):
            self.assertIn(phrase, agent)
        self.assertIn("必问项与条件触发项是硬门禁", readme)

    def test_wechat_builder_rejects_silent_auto_theme(self):
        builder = ROOT / "scripts/build_gzh_html.py"
        with tempfile.TemporaryDirectory() as tmp:
            directory = pathlib.Path(tmp)
            source = directory / "article-source.md"
            source.write_text(SAMPLE_MD, encoding="utf-8")
            state = self.write_task_state(directory)

            silent_auto = subprocess.run(
                [sys.executable, str(builder), str(source), "--no-preview", "--task-state", str(state)],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(silent_auto.returncode, 0)
            self.assertIn("不能静默使用主题 auto", silent_auto.stderr)

            confirmed_theme = subprocess.run(
                [sys.executable, str(builder), str(source), "--no-preview", "-t", "moyu-green", "--task-state", str(state)],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(confirmed_theme.returncode, 0)
            self.assertIn("不能由智能体静默指定主题", confirmed_theme.stderr)

            user_confirmed_theme = subprocess.run(
                [
                    sys.executable,
                    str(builder),
                    str(source),
                    "--no-preview",
                    "-t",
                    "moyu-green",
                    "--theme-confirmed",
                    "--task-state",
                    str(state),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(user_confirmed_theme.returncode, 0, user_confirmed_theme.stdout + user_confirmed_theme.stderr)

    def test_wechat_builder_rejects_silent_specific_theme(self):
        builder = ROOT / "scripts/build_html.py"
        with tempfile.TemporaryDirectory() as tmp:
            directory = pathlib.Path(tmp)
            source = directory / "article-source.md"
            output = directory / "article.html"
            source.write_text(SAMPLE_MD, encoding="utf-8")
            state = self.write_task_state(
                directory,
                delivery_style={
                    "value": "graphite-minimal",
                    "confirmed": True,
                    "source": "user",
                    "user_quote": "石墨极简风",
                },
            )

            silent_specific = subprocess.run(
                [
                    sys.executable,
                    str(builder),
                    str(source),
                    "-o",
                    str(output),
                    "--platform",
                    "公众号",
                    "-t",
                    "graphite-minimal",
                    "--task-state",
                    str(state),
                ],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(silent_specific.returncode, 0)
            self.assertIn("不能由智能体静默指定主题", silent_specific.stderr)

            confirmed_specific = subprocess.run(
                [
                    sys.executable,
                    str(builder),
                    str(source),
                    "-o",
                    str(output),
                    "--platform",
                    "公众号",
                    "-t",
                    "graphite-minimal",
                    "--theme-confirmed",
                    "--task-state",
                    str(state),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(confirmed_specific.returncode, 0, confirmed_specific.stdout + confirmed_specific.stderr)
            self.assertTrue(output.is_file())

    def test_wechat_copy_preview_writes_rich_html_clipboard(self):
        builder = ROOT / "scripts/build_gzh_html.py"
        with tempfile.TemporaryDirectory() as tmp:
            directory = pathlib.Path(tmp)
            source = directory / "article-source.md"
            output = directory / "article.html"
            source.write_text(TABLE_MD, encoding="utf-8")
            state = self.write_task_state(
                directory,
                delivery_style={"value": "red-white", "confirmed": True, "source": "user", "user_quote": "红白色系"},
            )

            result = subprocess.run(
                [
                    sys.executable,
                    str(builder),
                    str(source),
                    "-o",
                    str(output),
                    "-t",
                    "red-white",
                    "--theme-confirmed",
                    "--task-state",
                    str(state),
                ],
                capture_output=True,
                text=True,
            )
            preview = directory / "article-preview.html"
            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
            self.assertTrue(preview.is_file())
            text = preview.read_text(encoding="utf-8")
            self.assertIn("ClipboardItem", text)
            self.assertIn("'text/html'", text)
            self.assertIn("'text/plain'", text)
            self.assertIn("navigator.clipboard.write", text)
            self.assertIn("gzhFallbackCopy", text)
            self.assertIn("<meta charset=\"utf-8\">", text)

    def test_skill_requires_multiple_direction_options_with_ranking(self):
        skill = (ROOT / "SKILL.md").read_text(encoding="utf-8")
        readme = (ROOT / "README.md").read_text(encoding="utf-8")
        for text in (skill, readme):
            self.assertIn("3-5", text)
            self.assertIn("最推荐", text)
            self.assertIn("次推荐", text)
        self.assertIn("用户只确认平台和方向，不等于确认平台交付样式", skill)
        self.assertIn("交付样式未确认前，任务列表不得出现", skill)

    def test_three_layer_confirmation_policy_is_documented(self):
        skill = (ROOT / "SKILL.md").read_text(encoding="utf-8")
        platform = (ROOT / "references/platform-native-protocol.md").read_text(encoding="utf-8")
        readme = (ROOT / "README.md").read_text(encoding="utf-8")
        agent = (ROOT / "agents/openai.yaml").read_text(encoding="utf-8")

        for text in (skill, platform, readme):
            self.assertIn("必问", text)
            self.assertIn("条件问", text)
            self.assertIn("不默认问", text)
            for phrase in (
                "发布平台",
                "内容目标",
                "创作方向",
                "平台交付样式",
                "目标读者/阅读场景",
                "立场边界",
                "时效口径",
                "来源边界",
                "交付格式",
                "篇幅深度",
                "语气风格",
                "开头方式",
                "标题数量",
                "是否要金句",
                "是否要案例",
            ):
                self.assertIn(phrase, text)

        self.assertIn("three-layer confirmation", agent)
        self.assertIn("publishing platform, content goal", agent)
        self.assertIn("platform delivery style", agent)
        self.assertIn("do not ask by", agent)
        self.assertIn("default about tone", agent)

    def test_every_platform_has_delivery_style_confirmation(self):
        skill = (ROOT / "SKILL.md").read_text(encoding="utf-8")
        platform = (ROOT / "references/platform-native-protocol.md").read_text(encoding="utf-8")
        readme = (ROOT / "README.md").read_text(encoding="utf-8")
        delivery = (ROOT / "references/delivery-protocol.md").read_text(encoding="utf-8")

        for text in (skill, platform, readme):
            for phrase in (
                "平台交付样式",
                "回答 / 专栏",
                "手机卡片预览",
                "SEO-GEO 结构化",
                "静态 HTML",
            ):
                self.assertIn(phrase, text)

        for phrase in (
            "公众号确认排版主题",
            "知乎确认回答 / 专栏",
            "小红书确认纯文本笔记/手机卡片",
            "官网/网页确认网页结构",
            "个人博客确认 Markdown/CMS/静态 HTML",
        ):
            self.assertIn(phrase, delivery)

    def test_platform_builder_rejects_silent_wechat_auto_theme(self):
        builder = ROOT / "scripts/build_html.py"
        with tempfile.TemporaryDirectory() as tmp:
            directory = pathlib.Path(tmp)
            source = directory / "article-source.md"
            output = directory / "article.html"
            source.write_text(SAMPLE_MD, encoding="utf-8")
            auto_state = self.write_task_state(
                directory,
                delivery_style={
                    "value": "auto",
                    "confirmed": True,
                    "source": "auto_authorized",
                    "user_quote": "自动匹配",
                },
            )

            silent_auto = subprocess.run(
                [
                    sys.executable,
                    str(builder),
                    str(source),
                    "-o",
                    str(output),
                    "--platform",
                    "公众号",
                    "--task-state",
                    str(auto_state),
                ],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(silent_auto.returncode, 0)
            self.assertIn("不能静默使用主题 auto", silent_auto.stderr)

            authorized_auto = subprocess.run(
                [
                    sys.executable,
                    str(builder),
                    str(source),
                    "-o",
                    str(output),
                    "--platform",
                    "公众号",
                    "--auto-theme-ok",
                    "--task-state",
                    str(auto_state),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(authorized_auto.returncode, 0, authorized_auto.stdout + authorized_auto.stderr)
            self.assertTrue(output.is_file())

    def test_agent_entrypoint_routes_reader_value_before_novelty(self):
        text = (ROOT / "agents/openai.yaml").read_text(encoding="utf-8")
        self.assertIn("minimum necessary innovation", text)
        self.assertIn("reader job", text)
        self.assertIn("speaking position", text)

    def test_zhihu_and_xiaohongshu_require_layout_confirmation(self):
        skill = (ROOT / "SKILL.md").read_text(encoding="utf-8")
        platform = (ROOT / "references/platform-native-protocol.md").read_text(encoding="utf-8")
        agent = (ROOT / "agents/openai.yaml").read_text(encoding="utf-8")
        for text in (skill, platform):
            self.assertIn("是否需要 HTML 排版预览", text)
            self.assertIn("不要询问", text)
        self.assertIn("ask whether HTML layout preview is needed", agent)

    def test_platform_protocol_defines_minimum_necessary_confirmation_matrix(self):
        skill = (ROOT / "SKILL.md").read_text(encoding="utf-8")
        platform = (ROOT / "references/platform-native-protocol.md").read_text(encoding="utf-8")
        for phrase in (
            "最低必要询问",
            "平台交付样式",
            "公众号排版主题",
            "回答还是专栏",
            "内容目标",
            "发布环境",
            "Markdown / CMS 富文本 / 静态 HTML",
            "一次合并询问",
            "不重复询问",
        ):
            self.assertIn(phrase, platform)
        self.assertIn("必要询问矩阵", skill)
        self.assertIn("一次合并", skill)

    def test_non_layout_platform_bundle_requires_title_and_native_source_only(self):
        validator = ROOT / "scripts/validate_delivery_bundle.py"
        with tempfile.TemporaryDirectory() as tmp:
            delivery = pathlib.Path(tmp)
            state = self.write_task_state(
                delivery,
                platform={"value": "知乎", "confirmed": True, "source": "user", "user_quote": "知乎"},
                delivery_style={"value": "回答", "confirmed": True, "source": "user", "user_quote": "回答"},
            )
            (delivery / "title-strategy.md").write_text(
                "# 标题策略\n\n## 主标题\n测试标题\n\n## 备选标题\n备选一、备选二。",
                encoding="utf-8",
            )
            (delivery / "article-source.md").write_text("# 正文\n\n内容。", encoding="utf-8")
            result = subprocess.run(
                [sys.executable, str(validator), str(delivery), "--platform", "知乎", "--task-state", str(state)],
                capture_output=True,
                text=True,
            )
            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
            self.assertIn("平台原生正文", result.stdout)
            self.assertNotIn("跳过", result.stdout)

    def test_formatted_non_wechat_bundle_requires_clean_and_copy_html(self):
        validator = ROOT / "scripts/validate_delivery_bundle.py"
        with tempfile.TemporaryDirectory() as tmp:
            delivery = pathlib.Path(tmp)
            state = self.write_task_state(
                delivery,
                platform={"value": "知乎", "confirmed": True, "source": "user", "user_quote": "知乎"},
                delivery_style={"value": "回答 + HTML 预览", "confirmed": True, "source": "user", "user_quote": "回答 + HTML 预览"},
            )
            (delivery / "title-strategy.md").write_text(
                "# 标题策略\n\n## 主标题\n测试标题\n\n## 备选标题\n备选一、备选二。",
                encoding="utf-8",
            )
            (delivery / "article-source.md").write_text("# 正文\n\n内容。", encoding="utf-8")
            missing = subprocess.run(
                [sys.executable, str(validator), str(delivery), "--platform", "知乎", "--layout", "--task-state", str(state)],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(missing.returncode, 0)

            (delivery / "article.html").write_text("<main>正文</main>", encoding="utf-8")
            (delivery / "article-preview.html").write_text(
                "<button>复制正文</button><script>navigator.clipboard.writeText('正文')</script>",
                encoding="utf-8",
            )
            complete = subprocess.run(
                [sys.executable, str(validator), str(delivery), "--platform", "知乎", "--layout", "--task-state", str(state)],
                capture_output=True,
                text=True,
            )
            self.assertEqual(complete.returncode, 0, complete.stdout + complete.stderr)

    def test_xiaohongshu_bundle_accepts_plain_text_native_source(self):
        validator = ROOT / "scripts/validate_delivery_bundle.py"
        with tempfile.TemporaryDirectory() as tmp:
            delivery = pathlib.Path(tmp)
            state = self.write_task_state(
                delivery,
                platform={"value": "小红书", "confirmed": True, "source": "user", "user_quote": "小红书"},
                delivery_style={"value": "清爽纯文本笔记", "confirmed": True, "source": "user", "user_quote": "清爽纯文本笔记"},
            )
            (delivery / "title-strategy.md").write_text(
                "# 标题策略\n\n## 主标题\n测试标题\n\n## 备选标题\n备选一、备选二。",
                encoding="utf-8",
            )
            (delivery / "note-source.txt").write_text("测试笔记\n\n#话题一 #话题二", encoding="utf-8")
            result = subprocess.run(
                [sys.executable, str(validator), str(delivery), "--platform", "小红书", "--task-state", str(state)],
                capture_output=True,
                text=True,
            )
            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
            self.assertIn("平台原生正文", result.stdout)
            self.assertNotIn("跳过", result.stdout)

    def test_website_platform_requires_html_pair_by_default(self):
        validator = ROOT / "scripts/validate_delivery_bundle.py"
        with tempfile.TemporaryDirectory() as tmp:
            delivery = pathlib.Path(tmp)
            state = self.write_task_state(
                delivery,
                platform={"value": "官网/网页", "confirmed": True, "source": "user", "user_quote": "官网/网页"},
                delivery_style={"value": "普通网页文章", "confirmed": True, "source": "user", "user_quote": "普通网页文章"},
            )
            (delivery / "title-strategy.md").write_text(
                "# 标题策略\n\n## 主标题\n测试标题\n\n## 备选标题\n备选一、备选二。",
                encoding="utf-8",
            )
            (delivery / "web-source.md").write_text("# 网页正文\n\n内容。", encoding="utf-8")
            result = subprocess.run(
                [sys.executable, str(validator), str(delivery), "--platform", "官网/网页", "--task-state", str(state)],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("平台排版 HTML", result.stdout)
            self.assertIn("复制预览 HTML", result.stdout)

            (delivery / "web.html").write_text("<main>网页正文</main>", encoding="utf-8")
            (delivery / "web-preview.html").write_text(
                "<button>复制正文</button><script>function copyArticle(){}</script>",
                encoding="utf-8",
            )
            complete = subprocess.run(
                [sys.executable, str(validator), str(delivery), "--platform", "官网/网页", "--task-state", str(state)],
                capture_output=True,
                text=True,
            )
            self.assertEqual(complete.returncode, 0, complete.stdout + complete.stderr)
            for role in ("标题策略 Markdown", "平台原生正文", "平台排版 HTML", "复制预览 HTML"):
                self.assertIn(role, complete.stdout)

    def test_generic_builder_can_emit_clean_and_copy_preview_html(self):
        builder = ROOT / "scripts/build_html.py"
        with tempfile.TemporaryDirectory() as tmp:
            directory = pathlib.Path(tmp)
            source = directory / "article-source.md"
            output = directory / "article.html"
            source.write_text("# 测试标题\n\n正文内容。", encoding="utf-8")
            state = self.write_task_state(
                directory,
                platform={"value": "知乎", "confirmed": True, "source": "user", "user_quote": "知乎"},
                delivery_style={"value": "回答 + HTML 预览", "confirmed": True, "source": "user", "user_quote": "回答 + HTML 预览"},
            )
            result = subprocess.run(
                [
                    sys.executable,
                    str(builder),
                    str(source),
                    "-o",
                    str(output),
                    "--platform",
                    "知乎",
                    "--emit-pair",
                    "--task-state",
                    str(state),
                ],
                capture_output=True,
                text=True,
            )
            preview = directory / "article-preview.html"
            self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
            self.assertTrue(output.is_file())
            self.assertTrue(preview.is_file())
            self.assertNotIn("copyArticle", output.read_text(encoding="utf-8"))
            self.assertIn("copyArticle", preview.read_text(encoding="utf-8"))

    def test_xiaohongshu_copy_text_is_plain_publishable_text(self):
        source = "# 这是一条笔记\n\n**适合谁**\n\n- 新手\n\n## 标签\n\n#话题一 #话题二"
        plain = build_html.markdown_to_plain_text(source)
        self.assertNotIn("# 这是一条笔记", plain)
        self.assertNotIn("**适合谁**", plain)
        self.assertNotIn("## 标签", plain)
        self.assertIn("这是一条笔记", plain)
        self.assertIn("适合谁", plain)
        self.assertIn("#话题一 #话题二", plain)

    def test_generic_article_tables_wrap_instead_of_horizontal_scroll(self):
        self.assertIn("table-layout:fixed", build_html.HTML_TEMPLATE)
        self.assertIn("overflow-wrap:anywhere", build_html.HTML_TEMPLATE)
        self.assertNotIn(".article table{display:block;overflow-x:auto;}", build_html.HTML_TEMPLATE)

    def test_backup_gzh_preview_template_writes_rich_and_plain_clipboard(self):
        template = (ROOT / "assets/gzh-design/assets/preview-template.html").read_text(encoding="utf-8")
        self.assertIn("ClipboardItem", template)
        self.assertIn("'text/html'", template)
        self.assertIn("'text/plain'", template)

    def test_component_subheading_number_resets_per_chapter(self):
        markdown = """# 标题

## 第一章

### 第一个点

### 第二个点

## 第二章

### 新章节第一点
"""
        calls = []
        original = build_gzh.component_subheading

        def spy(theme_key, components, text, number):
            calls.append((text, number))
            return original(theme_key, components, text, number)

        with mock.patch.object(build_gzh, "component_subheading", side_effect=spy):
            build_gzh.build_component_section(markdown, "标题", "graphite-minimal")

        self.assertEqual(calls, [("第一个点", 1), ("第二个点", 2), ("新章节第一点", 1)])


class ReadmeDocumentationTests(unittest.TestCase):
    def test_markdown_table_separators_match_header_width(self):
        lines = (ROOT / "README.md").read_text(encoding="utf-8").splitlines()
        separator = re.compile(r"^\|(?:\s*:?-{3,}:?\s*\|)+$")
        for index, line in enumerate(lines):
            if separator.fullmatch(line):
                with self.subTest(line=index + 1):
                    self.assertGreater(index, 0)
                    self.assertEqual(lines[index - 1].count("|"), line.count("|"))

    def test_installation_is_runnable_from_the_public_repository(self):
        text = (ROOT / "README.md").read_text(encoding="utf-8")
        self.assertIn(
            "git clone https://github.com/NocodeMrLi/mr-li-writer-skill.git",
            text,
        )

    def test_directory_tree_mentions_delivery_protocol_and_tests(self):
        text = (ROOT / "README.md").read_text(encoding="utf-8")
        self.assertIn("delivery-protocol.md", text)
        self.assertIn("commercial-source-terms.txt", text)
        self.assertIn("├── tests/", text)
        self.assertIn("test_regressions.py", text)

    def test_promo_video_sources_are_trackable_but_outputs_stay_ignored(self):
        source = ROOT / "promo-video/scripts/generate-promo-video.js"
        output = ROOT / "promo-video/out/mr-li-writer-promo-silent.mp4"
        self.assertTrue(source.exists(), "宣传视频生成器源码缺失")
        source_check = subprocess.run(
            ["git", "check-ignore", "-q", str(source)],
            cwd=ROOT,
        )
        output_check = subprocess.run(
            ["git", "check-ignore", "-q", str(output)],
            cwd=ROOT,
        )
        self.assertNotEqual(source_check.returncode, 0)
        self.assertEqual(output_check.returncode, 0)

    def test_readme_documents_four_artifacts_and_runnable_checks(self):
        text = (ROOT / "README.md").read_text(encoding="utf-8")
        for artifact in ("title-strategy.md", "article-source.md", "article-gzh.html", "article-preview.html"):
            self.assertIn(artifact, text)
        self.assertIn("article-gzh-preview.html", text)
        self.assertIn("python3 -m unittest discover -s tests -v", text)
        self.assertIn("重要文章发布前", text)

    def test_readme_showcase_images_exist(self):
        text = (ROOT / "README.md").read_text(encoding="utf-8")
        paths = re.findall(r"\./(assets/readme-showcase/[^\"')> ]+\.jpg)", text)
        self.assertGreaterEqual(len(set(paths)), 12)
        for path in set(paths):
            self.assertTrue((ROOT / path).is_file(), path)

    def test_readme_long_showcase_images_share_top_aligned_crop_size(self):
        try:
            from PIL import Image
        except ImportError as exc:
            self.skipTest("Pillow unavailable: %s" % exc)
        sizes = {}
        for path in sorted((ROOT / "assets/readme-showcase").glob("*-long.jpg")):
            with Image.open(path) as image:
                sizes[path.name] = image.size
        self.assertEqual(len(sizes), 3)
        self.assertEqual(len(set(sizes.values())), 1, sizes)

    def test_readme_documents_adaptive_innovation_and_human_writing_inspiration(self):
        text = (ROOT / "README.md").read_text(encoding="utf-8")
        for phrase in ("最低必要创新", "读者任务", "点击契约", "说话位置"):
            self.assertIn(phrase, text)
        self.assertIn("https://github.com/KKKKhazix/human-writing", text)

    def test_public_notices_do_not_expose_local_source_paths(self):
        notice = (ROOT / "assets/gzh-design/NOTICE.md").read_text(encoding="utf-8")
        self.assertNotIn("/Users/", notice)

    def test_readme_explains_original_product_strengths_and_conditional_delivery(self):
        text = (ROOT / "README.md").read_text(encoding="utf-8")
        for phrase in ("不是拼装式提示词", "编辑路由", "最低必要创新", "写作到交付闭环"):
            self.assertIn(phrase, text)
        self.assertIn("需要排版", text)
        self.assertIn("不需要排版", text)
        self.assertIn("不机械凑四件套", text)

    def test_readme_explains_zhihu_xiaohongshu_layout_confirmation(self):
        text = (ROOT / "README.md").read_text(encoding="utf-8")
        self.assertIn("选择知乎或小红书后", text)
        self.assertIn("是否需要 HTML 排版预览", text)

    def test_readme_explains_platform_confirmation_points(self):
        text = (ROOT / "README.md").read_text(encoding="utf-8")
        self.assertIn("执行中会确认什么", text)
        for platform in ("公众号", "知乎", "小红书", "官网/网页", "个人博客"):
            self.assertIn(platform, text)
        self.assertIn("一次合并询问", text)


class ArticleLintTests(unittest.TestCase):
    def test_lint_warns_when_training_vendors_are_named_as_sources(self):
        article = """# 软考报名时间

数据来自全国软考报名平台当前页面及 51CTO、希赛网公开汇总。
"""
        with tempfile.TemporaryDirectory() as tmp:
            path = pathlib.Path(tmp) / "article.md"
            path.write_text(article, encoding="utf-8")
            output = io.StringIO()
            with mock.patch.object(sys, "argv", ["lint_article.py", str(path)]):
                with contextlib.redirect_stdout(output):
                    code = lint_article.main()
        self.assertEqual(code, 0)
        self.assertIn("商业相关第三方机构", output.getvalue())
        self.assertIn("据相关第三方机构公开信息/公开汇总", output.getvalue())
        self.assertIn("尚待官方确认", output.getvalue())

    def test_commercial_source_examples_are_configured_outside_scripts(self):
        terms = (ROOT / "references/commercial-source-terms.txt").read_text(encoding="utf-8")
        self.assertIn("51CTO", terms)
        self.assertIn("希赛网", terms)
        self.assertIn("才聚", terms)
        for script in ("scripts/lint_article.py", "scripts/build_html.py", "scripts/build_gzh_html.py"):
            source = (ROOT / script).read_text(encoding="utf-8")
            self.assertNotIn("51CTO", source)
            self.assertNotIn("希赛网", source)
            self.assertNotIn("才聚", source)

    def test_strict_lint_blocks_named_commercial_vendor_in_narrative_source_sentence(self):
        article = """# PMP 考点信息

才聚刚把 2026 年全部考点城市清单放出来：全国 51 城，覆盖七大区。
"""
        with tempfile.TemporaryDirectory() as tmp:
            path = pathlib.Path(tmp) / "article.md"
            path.write_text(article, encoding="utf-8")
            output = io.StringIO()
            with mock.patch.object(sys, "argv", ["lint_article.py", str(path), "--strict-delivery"]):
                with contextlib.redirect_stdout(output):
                    code = lint_article.main()
        self.assertEqual(code, 1)
        self.assertIn("商业相关第三方机构", output.getvalue())

    def test_strict_lint_uses_source_roles_instead_of_only_known_name_list(self):
        article = """# 考试城市信息

星河学习刚发布了全年考试城市汇总。
"""
        with tempfile.TemporaryDirectory() as tmp:
            directory = pathlib.Path(tmp)
            path = directory / "article.md"
            state = directory / "task-state.json"
            path.write_text(article, encoding="utf-8")
            state.write_text(
                json.dumps(
                    {
                        "research_scope": {
                            "source_entities": [
                                {
                                    "name": "星河学习",
                                    "role": "commercial_interested",
                                    "claim_scope": "考试城市汇总",
                                    "reader_visibility": "omit",
                                }
                            ]
                        }
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            output = io.StringIO()
            with mock.patch.object(
                sys,
                "argv",
                ["lint_article.py", str(path), "--strict-delivery", "--task-state", str(state)],
            ):
                with contextlib.redirect_stdout(output):
                    code = lint_article.main()
        self.assertEqual(code, 1)
        self.assertIn("商业相关第三方机构", output.getvalue())

    def test_html_builders_warn_on_named_commercial_vendor_without_source_keywords(self):
        articles = (
            "才聚刚把 2026 年全部考点城市清单放出来。",
            "启航培训机构发布了全年考试城市汇总。",
        )
        for article in articles:
            for warn in (build_html.warn_commercial_source_exposure, build_gzh.warn_commercial_source_exposure):
                with self.subTest(article=article, module=warn.__module__):
                    output = io.StringIO()
                    with contextlib.redirect_stderr(output):
                        warn(article)
                    self.assertIn("商业相关第三方机构", output.getvalue())

    def test_generic_commercial_patterns_do_not_treat_ministry_name_as_training_vendor(self):
        articles = (
            "教育部发布了最新通知。",
            "这座城市正在形成人才聚集效应。",
        )
        for article in articles:
            for warn in (build_html.warn_commercial_source_exposure, build_gzh.warn_commercial_source_exposure):
                with self.subTest(article=article, module=warn.__module__):
                    output = io.StringIO()
                    with contextlib.redirect_stderr(output):
                        warn(article)
                    self.assertNotIn("商业相关第三方机构", output.getvalue())

    def test_commercial_source_terms_missing_warns_before_generic_fallback(self):
        checks = (
            (build_html.configured_commercial_source_regex, "builtins.open"),
            (build_gzh.configured_commercial_source_regex, "builtins.open"),
            (lint_article.configured_commercial_source_patterns, "pathlib.Path.read_text"),
        )
        for func, patch_target in checks:
            with self.subTest(func=func.__module__):
                stderr = io.StringIO()
                with mock.patch(patch_target, side_effect=FileNotFoundError):
                    with contextlib.redirect_stderr(stderr):
                        result = func()
                self.assertIn("商业来源词表缺失", stderr.getvalue())
                self.assertFalse(result)

    def test_lint_accepts_generic_commercial_source_wording(self):
        article = """# 软考报名时间

数据来自全国软考报名平台当前页面及相关机构公开汇总。
"""
        with tempfile.TemporaryDirectory() as tmp:
            path = pathlib.Path(tmp) / "article.md"
            path.write_text(article, encoding="utf-8")
            output = io.StringIO()
            with mock.patch.object(sys, "argv", ["lint_article.py", str(path)]):
                with contextlib.redirect_stdout(output):
                    code = lint_article.main()
        self.assertEqual(code, 0)
        self.assertNotIn("商业相关第三方机构", output.getvalue())

    def test_lint_warns_on_reader_facing_recreation_disclaimer(self):
        article = """# 软考报名要求

正文。

本文为二创整理，具体政策以官方最新通知为准。
"""
        with tempfile.TemporaryDirectory() as tmp:
            path = pathlib.Path(tmp) / "article.md"
            path.write_text(article, encoding="utf-8")
            output = io.StringIO()
            with mock.patch.object(sys, "argv", ["lint_article.py", str(path)]):
                with contextlib.redirect_stdout(output):
                    code = lint_article.main()
        self.assertEqual(code, 0)
        self.assertIn("二创", output.getvalue())

    def test_impact_lint_does_not_require_formulaic_insight_markers(self):
        article = """# 普通工作日为什么更能检验健身计划

周二晚上六点，外面下着雨，人刚加完班。健身房还要坐四站地铁。

办卡只用一分钟，去一次却要换衣服、出门、训练、洗澡。计划能不能坚持，应该拿这样的普通一天来算。
"""
        with tempfile.TemporaryDirectory() as tmp:
            path = pathlib.Path(tmp) / "article.md"
            path.write_text(article, encoding="utf-8")
            output = io.StringIO()
            with mock.patch.object(
                sys,
                "argv",
                ["lint_article.py", str(path), "--impact-check", "--genre", "relationship-life"],
            ):
                with contextlib.redirect_stdout(output):
                    code = lint_article.main()
        self.assertEqual(code, 0)
        self.assertNotIn("未发现清晰的核心判断提示", output.getvalue())

    def test_lint_warns_on_semantic_reversal_posture(self):
        article = """# 为什么计划总会失败

你以为自己缺的是自律，其实真正的问题藏在时间安排里。

很多计划只写目标，没有给普通工作日留下执行空间。
"""
        with tempfile.TemporaryDirectory() as tmp:
            path = pathlib.Path(tmp) / "article.md"
            path.write_text(article, encoding="utf-8")
            output = io.StringIO()
            with mock.patch.object(sys, "argv", ["lint_article.py", str(path)]):
                with contextlib.redirect_stdout(output):
                    code = lint_article.main()
        self.assertEqual(code, 0)
        self.assertIn("翻案", output.getvalue())


class EditorialInnovationProtocolTests(unittest.TestCase):
    def test_skill_routes_editorial_value_before_innovation(self):
        text = (ROOT / "SKILL.md").read_text(encoding="utf-8")
        self.assertIn("editorial-routing-protocol.md", text)
        self.assertIn("最低必要创新", text)
        self.assertIn("读者任务", text)
        self.assertNotIn("保留完整创作流程：选题诊断、立意升级", text)

    def test_editorial_router_defines_four_innovation_levels(self):
        path = ROOT / "references/editorial-routing-protocol.md"
        self.assertTrue(path.exists())
        text = path.read_text(encoding="utf-8")
        for level in ("直给型", "微创新型", "视角转换型", "概念创新型"):
            self.assertIn(level, text)
        for reader_job in ("找答案", "做选择", "避风险", "被理解", "看故事"):
            self.assertIn(reader_job, text)
        self.assertIn("读者收益", text)
        self.assertIn("理解成本", text)

    def test_title_rules_use_click_contract_not_novelty_as_default(self):
        text = (ROOT / "references/title-rules.md").read_text(encoding="utf-8")
        self.assertIn("点击契约", text)
        self.assertIn("此刻相关性", text)
        self.assertIn("理解成本", text)
        self.assertIn("自然口语感", text)
        self.assertIn("熟悉的问题", text)
        self.assertIn("差异化不是必选项", text)

    def test_humanize_rules_define_positive_human_voice(self):
        text = (ROOT / "references/humanize-rules.md").read_text(encoding="utf-8")
        for phrase in ("说话位置", "选择性取舍", "每段新增", "承认不知道", "先写后审"):
            self.assertIn(phrase, text)

    def test_high_impact_protocol_has_optional_enhancement_levels(self):
        text = (ROOT / "references/high-impact-writing-protocol.md").read_text(encoding="utf-8")
        for level in ("基础发布", "编辑增强", "深度增强"):
            self.assertIn(level, text)
        self.assertIn("不需要升维", text)

    def test_editorial_benchmark_covers_twelve_cross_genre_cases(self):
        protocol = (ROOT / "references/editorial-evaluation-protocol.md").read_text(encoding="utf-8")
        cases = (ROOT / "tests/fixtures/editorial-routing-cases.md").read_text(encoding="utf-8")
        self.assertEqual(cases.count("## CASE-"), 12)
        for dimension in ("点击意愿", "理解成本", "可信度", "活人感", "平台原生度"):
            self.assertIn(dimension, protocol)
        self.assertIn("同题多跑", protocol)
        self.assertIn("盲评", protocol)


class GzhThemeRegressionTests(unittest.TestCase):
    LONG_QUOTE_MD = """# 软考含金量

> 含金量 = 证书权威性 × 用途契合度 × 环境认可度

## 第一部分：先判断证书

正文一。
"""
    LONG_PROSE_QUOTE_MD = """# AI 项目经理

> 学长说：这五项不需要你成为算法专家。需要的是你知道大模型能干什么、不能干什么、干起来要多久要多少钱。这个判断力，才是懂技术边界的真正含义。

## 第一部分：先判断边界

正文一。
"""

    def test_moyu_cover_renders_title_once_without_repeated_fragment(self):
        title = "一份不会把复杂问题写复杂的完整安装与避坑指南"
        rendered = build_gzh.build_component_section(SAMPLE_MD, title, "moyu-green")
        plain = html.unescape(re.sub(r"<[^>]+>", "", rendered))
        plain = re.sub(r"\s+", "", plain)
        self.assertEqual(plain.count(title), 1)

    def test_moyu_cover_separator_title_has_no_inline_highlight_gap(self):
        title = "9月第一周｜软考报名收官作战表"
        rendered = build_gzh.build_component_section(SAMPLE_MD, title, "moyu-green")
        plain = html.unescape(re.sub(r"<[^>]+>", "", rendered))
        plain = re.sub(r"[\s\u200b]+", "", plain)
        self.assertEqual(plain.count("9月第一周·软考报名收官作战表"), 1)
        self.assertNotIn('<span leaf="">｜软</span>', rendered)
        hero = rendered.split("03 PARTS", 1)[0]
        self.assertIn("9月第一周 · 软考报名收官作战表", hero.replace("\u200b", ""))
        self.assertIn("\u200b", hero)
        self.assertNotIn('<span leaf=""></span>', hero)
        self.assertNotRegex(hero, r"<p\b[^>]*>\s*</p>")
        self.assertNotIn("letter-spacing:-2px", hero)

    def test_separator_titles_use_connected_dot_across_cover_themes(self):
        title = "9月第一周｜软考报名收官作战表"
        for theme in ("moyu-green", "zen-whitespace", "moyu-ticket", "olive-journal"):
            with self.subTest(theme=theme):
                rendered = build_gzh.build_component_section(SAMPLE_MD, title, theme)
                self.assertIn("9月第一周 · 软考报名收官作战表", rendered.replace("\u200b", ""))
                self.assertIn("\u200b", rendered)
                self.assertNotIn(title, rendered)

    def test_cover_titles_prioritize_semantic_breaks_without_fixed_desktop_split(self):
        for theme in ("moyu-green", "zen-whitespace", "moyu-ticket", "olive-journal"):
            with self.subTest(theme=theme):
                rendered = build_gzh.build_component_section(
                    SAMPLE_MD,
                    "9月第一周｜软考报名收官作战表",
                    theme,
                )
                self.assertIn("word-break:keep-all", rendered)
                self.assertIn("overflow-wrap:anywhere", rendered)

    def test_olive_cover_keeps_title_and_illustration_side_by_side(self):
        rendered = build_gzh.build_component_section(
            SAMPLE_MD,
            "9月第一周｜软考报名收官作战表",
            "olive-journal",
        )
        self.assertIn("justify-content:center;gap:12px", rendered)
        self.assertIn("flex:1;min-width:0", rendered)
        self.assertIn("flex:0 0 22%;width:22%;min-width:64px;max-width:112px", rendered)
        hero = rendered.split("03 PARTS", 1)[0]
        self.assertNotIn("justify-content:center;gap:12px;flex-wrap:wrap", hero)

    def test_moyu_colored_title_uses_semantic_break_without_hanging_punctuation(self):
        title = "软考报名生死线：福建明天就截，缴费还比报名晚两天"
        rendered = build_gzh.build_component_section(SAMPLE_MD, title, "moyu-green")
        hero = rendered.split("03 PARTS", 1)[0]
        self.assertIn('<span leaf="">软考报名生死线</span>', hero)
        self.assertIn('<span leaf="">福建明天就截，缴费还比报名晚两天</span>', hero)
        self.assertNotIn('<span leaf="">软考报名生死线：', hero)
        self.assertNotRegex(hero, r'<span leaf="">[^<]*[，、；：:]</span>\s*</p>')

    def test_quote_cover_themes_render_the_intro_instead_of_an_empty_card(self):
        intro = "真正重要的不是多写一点，而是把值得说的话说清楚。"
        md = "# 标题\n\n%s\n\n## 第一部分\n正文。" % intro
        for theme in ("red-white", "graphite-minimal"):
            with self.subTest(theme=theme):
                rendered = build_gzh.component_hero(
                    theme,
                    build_gzh.load_theme_components(theme),
                    "测试文章",
                    build_gzh.parse_blocks(md),
                )
                plain = html.unescape(re.sub(r"<[^>]+>", "", rendered))
                plain = re.sub(r"\s+", "", plain)
                self.assertEqual(plain.count(intro), 1)
                self.assertNotRegex(rendered, r"<p\b[^>]*>\s*</p>")

    def test_graphite_section_number_is_a_placeholder(self):
        text = (ROOT / "assets/gzh-design/references/theme-graphite-minimal.md").read_text(encoding="utf-8")
        component = text.split("## 组件 5 章节标题", 1)[1].split("## 组件 6", 1)[0]
        self.assertIn("{{编号}}", component)
        self.assertNotIn('<span leaf="">01</span>', component)

    def test_every_theme_section_template_uses_dynamic_numbering(self):
        sections = {
            "theme-moyu-green.md": ("## 组件 4 章节标题", "## 组件 5"),
            "theme-red-white.md": ("## 组件 5 章节标题", "## 组件 6"),
            "theme-graphite-minimal.md": ("## 组件 5 章节标题", "## 组件 6"),
            "theme-zen-whitespace.md": ("## 组件 5 章节标题", "## 组件 6"),
            "theme-moyu-ticket.md": ("## 组件 3 章节标题", "## 组件 4"),
            "theme-olive-journal.md": ("## 组件 3 章节标题", "## 组件 4"),
        }
        theme_dir = ROOT / "assets/gzh-design/references"
        for filename, (start, end) in sections.items():
            with self.subTest(theme=filename):
                text = (theme_dir / filename).read_text(encoding="utf-8")
                component = text.split(start, 1)[1].split(end, 1)[0]
                self.assertIn("{{编号}}", component)

    def test_renderer_applies_each_section_number_to_every_theme(self):
        for theme in build_gzh.THEMES:
            components = build_gzh.load_theme_components(theme)
            for number in (1, 2, 3):
                with self.subTest(theme=theme, number=number):
                    rendered = build_gzh.component_section_title(theme, components, "章节标题", number)
                    plain = html.unescape(re.sub(r"<[^>]+>", "", rendered))
                    plain = re.sub(r"\s+", "", plain)
                    self.assertIn("%02d" % number, plain)
                    if number > 1:
                        self.assertNotIn("01", plain)

    def test_all_generated_themes_hide_production_labels(self):
        forbidden = ("公众号排版", "深度文章", "中文内容创作 Skill")
        for theme in build_gzh.THEMES:
            with self.subTest(theme=theme):
                rendered = build_gzh.build_component_section(SAMPLE_MD, "测试文章", theme)
                for phrase in forbidden:
                    self.assertNotIn(phrase, rendered)

    def test_moyu_toc_binds_each_heading_once(self):
        rendered = build_gzh.build_component_section(SAMPLE_MD, "测试文章", "moyu-green")
        for heading in ("第一部分：先判断问题", "第二部分：再给方法", "第三部分：完成交付"):
            self.assertEqual(rendered.count(heading), 2, heading)

    def test_moyu_toc_is_publication_safe_without_horizontal_scroll(self):
        rendered = build_gzh.build_component_section(SAMPLE_MD, "测试文章", "moyu-green")
        self.assertIn("overflow-wrap:anywhere", rendered)
        self.assertNotIn("overflow-x:auto", rendered)
        self.assertNotRegex(rendered, r"width:\s*\d+(?:\.\d+)?vw")
        self.assertNotIn("white-space:nowrap", rendered)
        self.assertNotIn("滑动查看", rendered)

    def test_theme_source_toc_components_are_publication_safe(self):
        theme_dir = ROOT / "assets/gzh-design/references"
        for path in sorted(theme_dir.glob("theme-*.md")):
            if path.name in {"theme-index.md", "theme-generator.md"}:
                continue
            text = path.read_text(encoding="utf-8")
            for match in re.finditer(
                r"^##\s+组件\s+\d+\s+([^\n]*(?:目录|导读|toc)[^\n]*)\n(.*?)(?=^##\s+组件\s+\d+|\Z)",
                text,
                re.I | re.M | re.S,
            ):
                with self.subTest(theme=path.name, component=match.group(1)):
                    component = match.group(2)
                    self.assertNotIn("overflow-x:auto", component)
                    self.assertNotRegex(component, r"width:\s*\d+(?:\.\d+)?vw")
                    self.assertNotIn("white-space:nowrap", component)

    def test_front_cover_badges_resist_vertical_wrapping_across_themes(self):
        for theme in ("moyu-green", "moyu-ticket", "olive-journal"):
            with self.subTest(theme=theme):
                rendered = build_gzh.build_component_section(SAMPLE_MD, "软考报名收官作战表", theme)
                self.assertIn("word-break:keep-all", rendered)
                self.assertIn("overflow-wrap:normal", rendered)
        for theme in ("moyu-green", "olive-journal"):
            with self.subTest(theme=theme):
                rendered = build_gzh.build_component_section(SAMPLE_MD, "软考报名收官作战表", theme)
                hero = rendered.split("03 PARTS", 1)[0]
                self.assertNotIn("justify-content:space-between;gap:", hero)
                self.assertNotIn("min-width:96px", hero)

    def test_markdown_tables_render_as_semantic_responsive_tables_in_every_theme(self):
        blocks = build_gzh.parse_blocks(TABLE_MD)
        self.assertTrue(any(block[0] == "table" for block in blocks))
        for theme in build_gzh.THEMES:
            with self.subTest(theme=theme):
                rendered = build_gzh.build_component_section(TABLE_MD, "报名状态", theme)
                self.assertIn("<table", rendered)
                self.assertIn("table-layout:fixed", rendered)
                self.assertIn("overflow-x:auto", rendered)
                self.assertIn("max-width:760px", rendered)
                self.assertRegex(rendered, r"min-width:(?:360|372|496|620|744|760)px")
                self.assertIn("word-break:keep-all", rendered)
                self.assertIn("overflow-wrap:anywhere", rendered)
                self.assertNotIn("| --- |", rendered)

    def test_generated_html_validator_rejects_raw_markdown_tables(self):
        validator = load_module("validate_gzh_html", "scripts/validate_gzh_html.py")
        source = '<section><p><span leaf="">| 省份 | 截止时间 |\n| --- | --- |</span></p></section>'
        errors, _, _ = validator.validate(source)
        self.assertTrue(any("Markdown 表格" in error for error in errors), errors)

    def test_generated_html_validator_allows_only_bounded_table_scrolling(self):
        validator = load_module("validate_gzh_html_table_scroll", "scripts/validate_gzh_html.py")
        allowed = (
            '<section style="overflow-x:auto;overflow-y:hidden;">'
            '<table style="width:100%;min-width:620px;max-width:760px;table-layout:fixed;">'
            '<tr><td><span leaf="">9月7日</span><br></td></tr></table></section>'
        )
        errors, _, _ = validator.validate(allowed)
        self.assertFalse(any("横向滚动" in error for error in errors), errors)

        for source in (
            '<section style="overflow-x:auto;"><p><span leaf="">目录卡片</span></p></section>',
            '<section style="overflow-x:auto;"><table style="min-width:1200px;table-layout:fixed;"><tr><td><span leaf="">数据</span></td></tr></table></section>',
        ):
            with self.subTest(source=source):
                errors, _, _ = validator.validate(source)
                self.assertTrue(any("横向滚动" in error for error in errors), errors)

    def test_generated_html_validator_reports_missing_file_without_traceback(self):
        validator = ROOT / "scripts/validate_gzh_html.py"
        result = subprocess.run(
            [sys.executable, str(validator), "/tmp/definitely-not-exist-mr-li.html"],
            capture_output=True,
            text=True,
        )
        self.assertEqual(result.returncode, 2)
        self.assertIn("[阻断] 无法读取 HTML 文件", result.stdout)
        self.assertNotIn("Traceback", result.stdout + result.stderr)

    def test_generated_html_validator_rejects_repeated_chapter_numbers(self):
        validator = load_module("validate_gzh_html_duplicate_chapters", "scripts/validate_gzh_html.py")
        source = """
<section>
  <section>
    <p><span leaf="">01</span></p><p><span leaf="">CHAPTER</span></p>
    <h3><span leaf="">第一章</span></h3>
  </section>
  <section>
    <p><span leaf="">01</span></p><p><span leaf="">CHAPTER</span></p>
    <h3><span leaf="">第二章</span></h3>
  </section>
</section>
"""
        errors, _, _ = validator.validate(source)
        self.assertTrue(any("章节编号" in error for error in errors), errors)

    def test_generated_html_validator_extracts_real_component_chapter_numbers(self):
        validator = load_module("validate_gzh_html_real_component_numbers", "scripts/validate_gzh_html.py")
        for theme in build_gzh.THEMES:
            with self.subTest(theme=theme):
                rendered = build_gzh.build_component_section(SAMPLE_MD, "测试文章", theme)
                self.assertEqual(validator.chapter_numbers(rendered), ["01", "02", "03"])

    def test_generated_html_validator_rejects_mutated_real_component_numbers(self):
        validator = load_module("validate_gzh_html_real_mutated_numbers", "scripts/validate_gzh_html.py")
        rendered = build_gzh.build_component_section(SAMPLE_MD, "测试文章", "moyu-green")
        rendered = rendered.replace("PART 02", "PART 04").replace("CHAPTER 02", "CHAPTER 04")
        errors, _, _ = validator.validate(rendered)
        self.assertTrue(any("编号不连续" in error for error in errors), errors)

    def test_generated_html_validator_rejects_repeated_real_body_chapter_numbers(self):
        validator = load_module("validate_gzh_html_real_repeated_body", "scripts/validate_gzh_html.py")
        rendered = build_gzh.build_component_section(SAMPLE_MD, "测试文章", "graphite-minimal")
        rendered = rendered.replace('<span leaf="">02</span>', '<span leaf="">01</span>')
        errors, _, _ = validator.validate(rendered)
        self.assertTrue(any("正文章节编号" in error for error in errors), errors)

    def test_generated_html_validator_reports_combined_marker_error_once(self):
        validator = load_module("validate_gzh_html_combined_dedupe", "scripts/validate_gzh_html.py")
        source = "".join(
            '<p><span leaf="">%02d · CHAPTER</span></p>' % number for number in (1, 2, 5)
        )
        errors, _, _ = validator.validate(source)
        numbered = [error for error in errors if "编号" in error]
        self.assertEqual(len(numbered), 1, errors)
        self.assertIn("01/02/05", numbered[0])

    def test_generated_html_validator_ignores_part_marker_inside_prose(self):
        validator = load_module("validate_gzh_html_prose_marker", "scripts/validate_gzh_html.py")
        source = '<section><p><span leaf="">正文提到 PART 03 的说法。</span></p></section>'
        errors, _, _ = validator.validate(source)
        self.assertFalse(any("编号" in error for error in errors), errors)

    def test_generated_html_validator_does_not_reject_id_equals_inside_reader_text(self):
        validator = load_module("validate_gzh_html_text_id", "scripts/validate_gzh_html.py")
        source = '<section><p><span leaf="">讨论 HTML 中 id=example 的写法。</span></p></section>'
        errors, _, _ = validator.validate(source)
        self.assertFalse(any("id 属性" in error for error in errors), errors)

    def test_generated_html_validator_rejects_real_id_attribute(self):
        validator = load_module("validate_gzh_html_real_id", "scripts/validate_gzh_html.py")
        source = '<section id="x"><p><span leaf="">正文。</span></p></section>'
        errors, _, _ = validator.validate(source)
        self.assertTrue(any("id 属性" in error for error in errors), errors)

    def test_wechat_builder_blocks_component_fallback_without_explicit_preview_flag(self):
        with tempfile.TemporaryDirectory() as tmp:
            directory = pathlib.Path(tmp)
            source = directory / "article-source.md"
            output = directory / "article.html"
            source.write_text(SAMPLE_MD, encoding="utf-8")
            with mock.patch.object(sys, "argv", [
                "build_gzh_html.py",
                str(source),
                "-o",
                str(output),
                "-t",
                "moyu-green",
                "--theme-confirmed",
                "--task-state",
                str(directory / "task-state.json"),
                "--no-preview",
            ]):
                with mock.patch.object(build_gzh, "require_task_state", return_value=0):
                    with mock.patch.object(build_gzh, "build_component_section", side_effect=RuntimeError("boom")):
                        stderr = io.StringIO()
                        with contextlib.redirect_stderr(stderr):
                            code = build_gzh.main()
            self.assertNotEqual(code, 0)
            self.assertIn("完整组件库渲染失败", stderr.getvalue())
            self.assertFalse(output.exists())

    def test_wechat_builder_allows_component_fallback_only_for_explicit_preview(self):
        with tempfile.TemporaryDirectory() as tmp:
            directory = pathlib.Path(tmp)
            source = directory / "article-source.md"
            output = directory / "article.html"
            source.write_text(SAMPLE_MD, encoding="utf-8")
            with mock.patch.object(sys, "argv", [
                "build_gzh_html.py",
                str(source),
                "-o",
                str(output),
                "-t",
                "moyu-green",
                "--theme-confirmed",
                "--task-state",
                str(directory / "task-state.json"),
                "--no-preview",
                "--allow-fallback-preview",
            ]):
                with mock.patch.object(build_gzh, "require_task_state", return_value=0):
                    with mock.patch.object(build_gzh, "validate_file", return_value=0):
                        with mock.patch.object(build_gzh, "build_component_section", side_effect=RuntimeError("boom")):
                            stderr = io.StringIO()
                            with contextlib.redirect_stderr(stderr):
                                code = build_gzh.main()
            self.assertEqual(code, 0, stderr.getvalue())
            self.assertIn("快速预览器", stderr.getvalue())
            self.assertTrue(output.exists())

    def test_fallback_signature_has_no_unfilled_placeholders(self):
        rendered = build_gzh.build_section("# 标题\n\n正文一段。", "标题", build_gzh.THEMES["moyu-green"])
        self.assertNotIn("{{", rendered)
        self.assertIn("我是 Mr.Li Writer", rendered)

    def test_multi_column_components_include_text_overflow_guards(self):
        checks = {
            "theme-graphite-minimal.md": ("## 组件 12 数据 / 要点卡片组", "## 组件 13"),
            "theme-red-white.md": ("## 组件 12 数据 / 要点卡片组", "## 组件 13"),
            "theme-moyu-green.md": ("## 组件 11 布局组件", "## 组件 12"),
            "theme-moyu-ticket.md": ("## 组件 2 票据封面", "## 组件 3"),
            "theme-olive-journal.md": ("## 组件 21 对比摘要卡", "## 组件 22"),
        }
        theme_dir = ROOT / "assets/gzh-design/references"
        for filename, (start, end) in checks.items():
            with self.subTest(theme=filename):
                text = (theme_dir / filename).read_text(encoding="utf-8")
                component = text.split(start, 1)[1].split(end, 1)[0]
                self.assertIn("min-width:0", component)
                self.assertIn("overflow-wrap:anywhere", component)

    def test_long_visual_quotes_are_balanced_before_wechat_wrapping(self):
        for theme in build_gzh.THEMES:
            with self.subTest(theme=theme):
                rendered = build_gzh.build_component_section(self.LONG_QUOTE_MD, "软考含金量", theme)
                self.assertIn('data-balanced="quote"', rendered)
                balanced = re.search(r'<section data-balanced="quote".*?</section>', rendered, re.S)
                self.assertIsNotNone(balanced)
                self.assertIn("<br>", balanced.group(0))

    def test_long_prose_quotes_are_left_aligned_not_centered(self):
        for theme in build_gzh.THEMES:
            with self.subTest(theme=theme):
                rendered = build_gzh.build_component_section(self.LONG_PROSE_QUOTE_MD, "AI 项目经理", theme)
                self.assertIn('data-balanced="prose-quote"', rendered)
                quote = re.search(r'<section data-balanced="prose-quote".*?</section>', rendered, re.S)
                self.assertIsNotNone(quote)
                self.assertIn("text-align:left", quote.group(0))
                self.assertNotIn("text-align:center", quote.group(0))

    def test_long_centered_emphasis_is_repaired_across_themes(self):
        sample = (
            '<p style="font-size:15px;margin:0 0 24px;text-align:center;color:#DC2626;'
            'font-weight:700;letter-spacing:1px;border-top:1px solid #FEE2E2;'
            'border-bottom:1px solid #FEE2E2;padding:14px 10px;">'
            '<span leaf="">报名截止 ≠ 缴费截止，但缴费截止才是真截止</span></p>'
        )
        repaired = build_gzh.repair_centered_text_orphans(sample)
        self.assertIn('data-balanced="center-text"', repaired)
        self.assertIn("text-align:left", repaired)
        self.assertNotIn("text-align:center", repaired)
        self.assertIn("overflow-wrap:anywhere", repaired)

    def test_short_centered_emphasis_can_stay_centered(self):
        sample = (
            '<p style="font-size:15px;text-align:center;font-weight:700;">'
            '<span leaf="">真正重要</span></p>'
        )
        repaired = build_gzh.repair_centered_text_orphans(sample)
        self.assertIn("text-align:center", repaired)
        self.assertNotIn('data-balanced="center-text"', repaired)

    def test_public_topic_tags_do_not_expose_editorial_posture_words(self):
        labels = build_gzh.public_topic_tags("软考含金量中立深度解读", [])
        self.assertNotIn("中立", labels)
        self.assertEqual(labels, ("深度解读", "判断参考"))

    def test_front_labels_do_not_expose_information_demystification_capability(self):
        for internal_label in ("信息祛魅", "信息去魅"):
            with self.subTest(label=internal_label):
                self.assertEqual(build_gzh.safe_front_label(internal_label, "信息指南"), "信息指南")

    def test_front_labels_use_allowlist_not_an_expanding_internal_word_blacklist(self):
        for internal_label in ("内容降噪", "事实筛选", "来源治理", "编辑增强", "反营销审查"):
            with self.subTest(label=internal_label):
                self.assertEqual(build_gzh.safe_front_label(internal_label, "信息指南"), "信息指南")
        for public_label in ("最新消息", "考试动态", "政策动态", "前沿观察"):
            with self.subTest(label=public_label):
                self.assertEqual(build_gzh.safe_front_label(public_label), public_label)

    def test_validator_rejects_forbidden_reader_facing_badges(self):
        validator = load_module("validate_gzh_html_front_badge", "scripts/validate_gzh_html.py")
        for label in ("中立", "信息祛魅", "信息去魅"):
            with self.subTest(label=label):
                source = '<section><span style="padding:2px 8px;border-radius:4px;"><span leaf="">%s</span></span></section>' % label
                errors, _, _ = validator.validate(source)
                self.assertTrue(any("前端标签" in error for error in errors), errors)

    def test_skill_documents_frontend_badge_and_title_linebreak_guards(self):
        skill = (ROOT / "SKILL.md").read_text(encoding="utf-8")
        protocol = (ROOT / "references/delivery-protocol.md").read_text(encoding="utf-8")
        for text in (skill, protocol):
            self.assertIn("前端标签白名单", text)
            self.assertIn("标题/金句换行", text)


if __name__ == "__main__":
    unittest.main()
